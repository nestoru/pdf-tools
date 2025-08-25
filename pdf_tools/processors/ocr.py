# pdf_tools/processors/ocr.py

import logging
import os
import sys
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Optional, List

import ocrmypdf
from pydantic import BaseModel, Field, validator

# Configure logging
logging.basicConfig(
    format='[%(asctime)s] %(levelname)-8s %(message)s',
    datefmt='%m/%d/%y %H:%M:%S',
    level=logging.INFO
)

logger = logging.getLogger(__name__)

class OCRConfig(BaseModel):
    """Configuration for OCR processing."""

    languages: list[str] = Field(
        default_factory=lambda: ['eng'],
        description="Languages to use for OCR (ISO 639-2 codes)"
    )
    dpi: int = Field(
        default=300,
        ge=72,
        le=600,
        description="DPI for PDF to image conversion"
    )
    force_ocr: bool = Field(
        default=False,
        description="Force OCR on every page, even if text is detected"
    )
    skip_text: bool = Field(
        default=False,
        description="Skip OCR on pages that already contain text"
    )
    rotate_pages: bool = Field(
        default=True,
        description="Automatically rotate pages based on detected text orientation"
    )
    clean: bool = Field(
        default=True,
        description="Clean visual noise from scanned images"
    )
    output_formats: list[str] = Field(
        default_factory=lambda: ['pdf'],
        description="Output formats to generate (pdf, docx)"
    )

    @validator('languages')
    def validate_languages(cls, v):
        """Validate that at least one language is provided and they are ISO 639-2."""
        if not v:
            raise ValueError("At least one language must be specified")

        # Convert any 2-letter codes to 3-letter codes
        lang_map = {'en': 'eng', 'fr': 'fra', 'de': 'deu', 'es': 'spa'}
        return [lang_map.get(lang, lang) for lang in v]

    @validator('output_formats')
    def validate_output_formats(cls, v):
        """Validate output formats."""
        if not v:
            raise ValueError("At least one output format must be specified")
        
        valid_formats = {'pdf', 'docx'}
        invalid_formats = set(v) - valid_formats
        if invalid_formats:
            raise ValueError(f"Invalid output formats: {', '.join(invalid_formats)}. Valid formats: {', '.join(valid_formats)}")
        
        return v

class OCRResult(BaseModel):
    """Result of OCR processing."""

    success: bool = Field(..., description="Whether processing was successful")
    error_message: Optional[str] = Field(None, description="Error message if processing failed")
    pages_processed: int = Field(0, ge=0, description="Number of pages processed")
    processing_time: float = Field(..., description="Processing time in seconds")
    pdf_generated: bool = Field(default=False, description="Whether PDF was generated")
    docx_generated: bool = Field(default=False, description="Whether DOCX was generated")

class DirectoryProcessResult(BaseModel):
    """Result of processing a directory of PDFs."""

    processed_count: int = Field(0, ge=0, description="Number of files processed")
    skipped_count: int = Field(0, ge=0, description="Number of files skipped")
    errors: list[str] = Field(default_factory=list, description="List of error messages")
    total_processing_time: float = Field(..., description="Total processing time in seconds")
    docx_generated: int = Field(0, ge=0, description="Number of DOCX files generated")

class PDFOCRProcessor:
    """Handles OCR processing for PDF files."""

    def __init__(self, config: Optional[OCRConfig] = None):
        """Initialize the OCR processor."""
        self.config = config or OCRConfig()
        logger.debug(f"Initialized with languages: {self.config.languages}")
        logger.debug(f"Output formats: {self.config.output_formats}")

    def should_process(self, input_path: Path, output_dir: Path, relative_path: Path) -> bool:
        """Check if the PDF needs to be processed based on timestamps and requested formats."""
        
        # Check each requested output format
        pdf_output_path = output_dir / relative_path.parent / f"{relative_path.stem}.pdf"
        docx_output_path = output_dir / relative_path.parent / f"{relative_path.stem}.docx"
        
        needs_processing = False
        
        # Check PDF if requested
        if 'pdf' in self.config.output_formats:
            if not pdf_output_path.exists():
                needs_processing = True
                logger.debug(f"PDF output missing: {pdf_output_path}")
            elif pdf_output_path.stat().st_size == 0:
                logger.warning(f"PDF output exists but is empty: {pdf_output_path}")
                needs_processing = True
            else:
                # Check timestamps
                input_time = datetime.fromtimestamp(input_path.stat().st_mtime)
                output_time = datetime.fromtimestamp(pdf_output_path.stat().st_mtime)
                if input_time > output_time:
                    needs_processing = True
                    logger.debug(f"PDF output older than input: {pdf_output_path}")
        
        # Check DOCX if requested
        if 'docx' in self.config.output_formats:
            if not docx_output_path.exists():
                needs_processing = True
                logger.debug(f"DOCX output missing: {docx_output_path}")
            elif docx_output_path.stat().st_size == 0:
                logger.warning(f"DOCX output exists but is empty: {docx_output_path}")
                needs_processing = True
            else:
                # Check timestamps - compare against the PDF if it exists, otherwise against input
                source_path = pdf_output_path if pdf_output_path.exists() else input_path
                source_time = datetime.fromtimestamp(source_path.stat().st_mtime)
                output_time = datetime.fromtimestamp(docx_output_path.stat().st_mtime)
                if source_time > output_time:
                    needs_processing = True
                    logger.debug(f"DOCX output older than source: {docx_output_path}")
        
        return needs_processing

    def check_pdf2docx_available(self) -> bool:
        """Check if pdf2docx library is available for DOCX conversion."""
        try:
            import pdf2docx
            logger.info("pdf2docx library found and available")
            return True
        except ImportError:
            logger.error("pdf2docx library not found")
            return False

    def convert_pdf_to_docx(self, pdf_path: Path, docx_path: Path) -> bool:
        """Convert PDF to DOCX by extracting OCR'd text and recreating layout."""
        try:
            logger.info(f"Converting PDF to DOCX with OCR text extraction: {pdf_path} -> {docx_path}")
            
            # First, extract the actual text from the OCR'd PDF
            extracted_text = self._extract_ocr_text_from_pdf(pdf_path)
            if not extracted_text:
                logger.error("No text could be extracted from the PDF")
                return False
            
            # Create DOCX with the extracted text
            return self._create_docx_from_text(extracted_text, docx_path)
            
        except Exception as e:
            logger.error(f"Error converting PDF to DOCX: {str(e)}")
            # Clean up empty DOCX file if it was created
            if docx_path.exists() and docx_path.stat().st_size == 0:
                try:
                    docx_path.unlink()
                    logger.info(f"Removed empty DOCX file: {docx_path}")
                except Exception:
                    pass
            return False

    def _extract_ocr_text_from_pdf(self, pdf_path: Path) -> List[dict]:
        """Extract OCR'd text from PDF using PyMuPDF, which can read invisible text layers."""
        try:
            import fitz  # PyMuPDF
            
            logger.info(f"Extracting OCR'd text from PDF: {pdf_path}")
            
            doc = fitz.open(pdf_path)
            pages_text = []
            
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                
                # Extract text with position and formatting information
                # This method can extract invisible OCR text that other tools miss
                text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE | fitz.TEXT_PRESERVE_LIGATURES)
                
                # Also try to get plain text as fallback
                plain_text = page.get_text()
                
                page_info = {
                    'page_number': page_num + 1,
                    'width': page.rect.width,
                    'height': page.rect.height,
                    'text_dict': text_dict,
                    'plain_text': plain_text.strip()
                }
                
                pages_text.append(page_info)
                
                # Log what we found
                char_count = len(plain_text.strip())
                blocks_count = len(text_dict.get('blocks', []))
                logger.info(f"Page {page_num + 1}: {char_count} characters, {blocks_count} blocks")
            
            doc.close()
            
            total_chars = sum(len(page['plain_text']) for page in pages_text)
            logger.info(f"Total extracted text: {total_chars} characters across {len(pages_text)} pages")
            
            if total_chars == 0:
                logger.warning("No text was extracted from the PDF - the OCR may have failed or the PDF may not contain text")
                return []
            
            return pages_text
            
        except ImportError:
            logger.error("PyMuPDF (fitz) library not found. Install with: pip install PyMuPDF")
            return []
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            return []

    def _create_docx_from_text(self, pages_text: List[dict], docx_path: Path) -> bool:
        """Create a DOCX document from extracted text while preserving basic layout."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            logger.info(f"Creating DOCX from extracted text: {docx_path}")
            
            # Ensure output directory exists
            docx_path.parent.mkdir(parents=True, exist_ok=True)
            
            doc = Document()
            
            # Set narrow margins for better space utilization
            for section in doc.sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            for page_info in pages_text:
                if page_info['page_number'] > 1:
                    # Add page break between pages
                    doc.add_page_break()
                
                # Add page header
                page_header = doc.add_paragraph()
                page_header.add_run(f"Page {page_info['page_number']}").bold = True
                page_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Process the text blocks to maintain some structure
                self._add_text_blocks_to_docx(doc, page_info)
            
            doc.save(docx_path)
            
            # Verify the output
            file_size = docx_path.stat().st_size
            logger.info(f"Successfully created DOCX: {docx_path} (size: {file_size} bytes)")
            
            # Verify content
            self._verify_docx_content(docx_path)
            
            return True
            
        except ImportError as e:
            logger.error(f"Required library not found for DOCX creation: {str(e)}")
            logger.error("Install with: pip install python-docx PyMuPDF")
            return False
        except Exception as e:
            logger.error(f"Error creating DOCX: {str(e)}")
            return False

    def _add_text_blocks_to_docx(self, doc, page_info):
        """Add text blocks from a page to the DOCX document."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # First, try to use structured text from text_dict
        text_dict = page_info.get('text_dict', {})
        blocks = text_dict.get('blocks', [])
        
        if blocks:
            # Process structured text blocks
            for block in blocks:
                if block.get('type') == 0:  # Text block
                    self._process_text_block(doc, block)
                elif block.get('type') == 1:  # Image block
                    # Add image placeholder
                    p = doc.add_paragraph()
                    run = p.add_run(f"[IMAGE: {block.get('width', 0):.0f}x{block.get('height', 0):.0f}]")
                    run.italic = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Fallback to plain text if structured data is not available
            plain_text = page_info.get('plain_text', '')
            if plain_text:
                # Split into paragraphs and add them
                paragraphs = plain_text.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        p = doc.add_paragraph(para_text.strip())

    def _process_text_block(self, doc, block):
        """Process a single text block from the PDF."""
        from docx.shared import Pt  # Import here to avoid issues
        
        lines = block.get('lines', [])
        if not lines:
            return
        
        # Create a paragraph for this block
        p = doc.add_paragraph()
        
        for line in lines:
            spans = line.get('spans', [])
            for span in spans:
                text = span.get('text', '')
                if text.strip():
                    run = p.add_run(text)
                    
                    # Apply formatting if available
                    font_size = span.get('size', 12)
                    if font_size and font_size > 0:
                        run.font.size = Pt(min(font_size, 18))  # Cap font size
                    
                    font_flags = span.get('flags', 0)
                    if font_flags & 16:  # Bold
                        run.bold = True
                    if font_flags & 2:   # Italic
                        run.italic = True
            
            # Add line break if this isn't the last line
            if line != lines[-1]:
                p.add_run('\n')

    def _verify_docx_content(self, docx_path: Path) -> None:
        """Verify that the DOCX file contains actual text content, not just images."""
        try:
            from docx import Document
            
            doc = Document(docx_path)
            text_content = []
            image_count = 0
            
            # Count text and images in the document
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Count images in the document
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            
            text_chars = sum(len(text) for text in text_content)
            
            logger.info(f"DOCX content verification: {len(text_content)} paragraphs, "
                       f"{text_chars} text characters, {image_count} images")
            
            if text_chars == 0 and image_count > 0:
                logger.warning("DOCX appears to contain only images - OCR text may not have been extracted properly")
            elif text_chars > 0:
                logger.info(f"DOCX contains extracted text: {text_chars} characters")
                # Show a sample of the extracted text for verification
                if text_content:
                    sample = text_content[0][:100] + "..." if len(text_content[0]) > 100 else text_content[0]
                    logger.info(f"Sample text: {sample}")
            
        except ImportError:
            logger.warning("python-docx not available for content verification")
        except Exception as e:
            logger.warning(f"Could not verify DOCX content: {str(e)}")

    def process_pdf(self, input_path: Path, output_dir: Path, relative_path: Path) -> OCRResult:
        """Process a single PDF file with OCR and save the result(s)."""
        start_time = datetime.now()
        pdf_generated = False
        docx_generated = False

        try:
            logger.info(f"Processing file: {input_path.absolute()}")

            # Ensure output directory exists
            output_dir.mkdir(parents=True, exist_ok=True)

            # Generate PDF output path
            pdf_output_path = output_dir / relative_path.parent / f"{relative_path.stem}.pdf"
            pdf_output_path.parent.mkdir(parents=True, exist_ok=True)

            # Process PDF with OCR if PDF format is requested
            if 'pdf' in self.config.output_formats:
                try:
                    result = ocrmypdf.ocr(
                        input_file=str(input_path),
                        output_file=str(pdf_output_path),
                        language='+'.join(self.config.languages),
                        deskew=True,
                        clean=self.config.clean,
                        rotate_pages=self.config.rotate_pages,
                        force_ocr=self.config.force_ocr,
                        skip_text=self.config.skip_text,
                        progress_bar=False,
                        optimize=1,
                        output_type='pdf',
                        use_threads=True,
                        jobs=os.cpu_count()
                    )
                except Exception as e:
                    # If standard approach fails with image size error, try command line approach
                    if "Image size" in str(e) and "pixels exceeds limit" in str(e):
                        logger.warning(f"Image size limit exceeded, trying command line approach")

                        cmd = [
                            "ocrmypdf",
                            "--language", "+".join(self.config.languages),
                            "--deskew",
                            "--optimize", "1",
                        ]

                        if self.config.clean:
                            cmd.append("--clean")
                        if self.config.rotate_pages:
                            cmd.append("--rotate-pages")
                        if self.config.force_ocr:
                            cmd.append("--force-ocr")
                        if self.config.skip_text:
                            cmd.append("--skip-text")

                        cmd.extend([str(input_path), str(pdf_output_path)])

                        logger.info(f"Running command: {' '.join(cmd)}")

                        try:
                            process = subprocess.run(
                                cmd,
                                env={
                                    "OCRMYPDF_PIKEPDF_IMAGE_MAX_PIXELS": "2000000000",  # 2 billion pixels
                                    "OCRMYPDF_IMAGE_MAX_PIXELS": "2000000000",  # Ensure this is set for any possible overrides
                                },
                                check=True,
                                capture_output=True,
                                text=True
                            )
                            result = 0  # Success
                        except subprocess.CalledProcessError as e:
                            raise Exception(f"Command line OCR failed: {str(e)}")
                    else:
                        # Re-raise other errors
                        raise

                if result == 0:  # Success
                    # Verify the output file exists and has content
                    if not pdf_output_path.exists():
                        raise Exception(f"OCR completed but PDF output file does not exist")

                    if pdf_output_path.stat().st_size == 0:
                        raise Exception(f"OCR completed but PDF output file is empty")

                    pdf_generated = True
                    logger.info(f"Successfully generated PDF: {pdf_output_path}")
                else:
                    error_msg = f"OCR failed with exit code: {result}"
                    logger.error(error_msg)
                    return OCRResult(
                        success=False,
                        error_message=error_msg,
                        pages_processed=0,
                        processing_time=(datetime.now() - start_time).total_seconds(),
                        pdf_generated=False,
                        docx_generated=False
                    )

            # Convert to DOCX if requested
            if 'docx' in self.config.output_formats:
                # Use the OCR-ed PDF as source if it exists, otherwise use the original
                source_pdf = pdf_output_path if pdf_generated else input_path
                docx_output_path = output_dir / relative_path.parent / f"{relative_path.stem}.docx"
                
                # pdf2docx is required for DOCX conversion
                if self.convert_pdf_to_docx(source_pdf, docx_output_path):
                    docx_generated = True
                else:
                    # If pdf2docx conversion fails, this is a hard error
                    error_msg = "pdf2docx DOCX conversion failed"
                    logger.error(error_msg)
                    return OCRResult(
                        success=False,
                        error_message=error_msg,
                        pages_processed=0,
                        processing_time=(datetime.now() - start_time).total_seconds(),
                        pdf_generated=pdf_generated,
                        docx_generated=False
                    )

            # Get page count from processed or original PDF
            try:
                if pdf_generated:
                    with open(pdf_output_path, 'rb') as f:
                        from pypdf import PdfReader
                        pdf = PdfReader(f)
                        num_pages = len(pdf.pages)
                else:
                    with open(input_path, 'rb') as f:
                        from pypdf import PdfReader
                        pdf = PdfReader(f)
                        num_pages = len(pdf.pages)
            except Exception as e:
                logger.warning(f"Could not get page count: {str(e)}")
                num_pages = 0

            # Determine overall success
            requested_pdf = 'pdf' in self.config.output_formats
            requested_docx = 'docx' in self.config.output_formats
            
            success = True
            if requested_pdf and not pdf_generated:
                success = False
            if requested_docx and not docx_generated:
                success = False

            if success:
                formats_generated = []
                if pdf_generated:
                    formats_generated.append("PDF")
                if docx_generated:
                    formats_generated.append("DOCX")
                logger.info(f"Successfully processed file: {input_path.absolute()} -> {', '.join(formats_generated)}")
            
            return OCRResult(
                success=success,
                pages_processed=num_pages,
                processing_time=(datetime.now() - start_time).total_seconds(),
                pdf_generated=pdf_generated,
                docx_generated=docx_generated
            )

        except Exception as e:
            error_msg = f"Error processing {input_path.name}: {str(e)}"
            logger.error(error_msg)

            # If an error occurred, make sure to clean up any empty output files
            if 'pdf' in self.config.output_formats:
                pdf_output_path = output_dir / relative_path.parent / f"{relative_path.stem}.pdf"
                if pdf_output_path.exists() and pdf_output_path.stat().st_size == 0:
                    try:
                        pdf_output_path.unlink()
                        logger.info(f"Removed empty PDF file after error: {pdf_output_path.absolute()}")
                    except Exception:
                        pass
            
            if 'docx' in self.config.output_formats:
                docx_output_path = output_dir / relative_path.parent / f"{relative_path.stem}.docx"
                if docx_output_path.exists() and docx_output_path.stat().st_size == 0:
                    try:
                        docx_output_path.unlink()
                        logger.info(f"Removed empty DOCX file after error: {docx_output_path.absolute()}")
                    except Exception:
                        pass

            return OCRResult(
                success=False,
                error_message=error_msg,
                pages_processed=0,
                processing_time=(datetime.now() - start_time).total_seconds(),
                pdf_generated=False,
                docx_generated=False
            )

    def find_pdf_files(self, directory: Path) -> List[Path]:
        """Find all PDF files in a directory, case-insensitive for extension."""
        pdf_files = []
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() == '.pdf':
                pdf_files.append(file_path)
        return pdf_files

    def process_directory(self, input_dir: Path, output_dir: Path) -> DirectoryProcessResult:
        """Process all PDF files in a directory, including subdirectories."""
        if not input_dir.is_dir():
            raise ValueError(f"Input directory does not exist: {input_dir}")

        output_dir.mkdir(parents=True, exist_ok=True)

        start_time = datetime.now()
        processed_count = 0
        skipped_count = 0
        docx_generated_count = 0
        errors = []

        # Recursively process all PDF files
        for input_path in self.find_pdf_files(input_dir):
            # Skip input files with zero size
            if input_path.stat().st_size == 0:
                logger.warning(f"Skipping empty input file: {input_path.absolute()}")
                continue

            # Calculate the relative path
            relative_path = input_path.relative_to(input_dir)

            # Check if we should process this file
            if self.should_process(input_path, output_dir, relative_path):
                result = self.process_pdf(input_path, output_dir, relative_path)
                if result.success:
                    processed_count += 1
                    if result.docx_generated:
                        docx_generated_count += 1
                else:
                    errors.append(result.error_message)
            else:
                # Determine what files already exist to provide better logging
                pdf_output_path = output_dir / relative_path.parent / f"{relative_path.stem}.pdf"
                docx_output_path = output_dir / relative_path.parent / f"{relative_path.stem}.docx"
                
                existing_files = []
                if pdf_output_path.exists() and pdf_output_path.stat().st_size > 0:
                    existing_files.append("PDF")
                if docx_output_path.exists() and docx_output_path.stat().st_size > 0:
                    existing_files.append("DOCX")
                
                if existing_files:
                    logger.info(f"Skipping {input_path.name} - {', '.join(existing_files)} already up to date")
                else:
                    logger.info(f"Skipping {input_path.name} - outputs are newer than input")
                skipped_count += 1

        return DirectoryProcessResult(
            processed_count=processed_count,
            skipped_count=skipped_count,
            errors=errors,
            total_processing_time=(datetime.now() - start_time).total_seconds(),
            docx_generated=docx_generated_count
        )
